#coding:utf-8

import StringIO

from	reportlab.pdfgen	import	canvas
from	reportlab.lib.units	import	mm
from	reportlab.pdfbase	import	pdfmetrics
from	reportlab.pdfbase	import	ttfonts
from	reportlab.lib		import	colors
from	reportlab.lib.pagesizes	import	letter, A4, landscape

from	reportlab.platypus.tables	import	Table, TableStyle
from	reportlab.platypus.doctemplate	import	SimpleDocTemplate
from	reportlab.platypus.paragraph	import	Paragraph
from	reportlab.lib.styles		import	ParagraphStyle,getSampleStyleSheet
from	reportlab.platypus		import	Frame,Spacer

from	reportlab.platypus		import	Image


from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE,WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml



from	kis.lib.contract	import	FIO_Job_Person, FIO_Job_Boss, GetSingD, GetSingP



### --- pdf file для штампа согласования ---
def	Stamp(buff,d_id,person):
    
    ### --- Список согласователей ---
    per = FIO_Job_Person(person)

    Font1 = ttfonts.TTFont('PT','kis/fonts/PTC55F.ttf')
    Font2 = ttfonts.TTFont('PTB','kis/fonts/PTC75F.ttf')
    Font3 = ttfonts.TTFont('PTI','kis/fonts/PTS56F.ttf')

    pdfmetrics.registerFont(Font1)
    pdfmetrics.registerFont(Font2)
    pdfmetrics.registerFont(Font3)


    style = getSampleStyleSheet()
    style.add(ParagraphStyle(name='Head',wordWrap=True,fontName='PTB',fontSize=14,spaceAfter=5*mm,spaceBefore=5*mm,alignment=1))
    style.add(ParagraphStyle(name='DepName',wordWrap=True,fontName='PTB',fontSize=10,spaceAfter=5*mm,spaceBefore=5*mm,alignment=1))
    style.add(ParagraphStyle(name='Data',wordWrap=True,fontName='PT',fontSize=8,spaceAfter=1*mm,spaceBefore=1*mm,alignment=0))
    
    doc = SimpleDocTemplate(buff,topMargin=10*mm,bottomMargin=10*mm,leftMargin=20*mm,rightMargin=10*mm)

    elements = []

    elements.append(Paragraph('КИС Договоры заявки',style["Head"]))
    elements.append(Paragraph('Номер заявки '+str(d_id),style["Head"]))


    Tdata = [['Участники договорной\nработы','ФИО лица,\nзавизировавшего договор'],]

    author = per[0]
    try:
	Tdata.append([Paragraph(u'Ответственный исполнитель, '+author[0],style["Data"]),Paragraph(author[1],style["Data"])],)
    except:
	pass

    for item in per[1:]:
	Tdata.append([Paragraph(item[0],style["Data"]),Paragraph(item[1],style["Data"])],)
    
    TableHead=Table(Tdata)
    TableHead.setStyle([('FONTNAME',(0,0),(-1,-1),'PTB'),
		('FONTSIZE',(0,0),(-1,-1),10),
		('ALIGN',(0,0),(-1,0),'CENTER'),
		('ALIGN',(0,1),(-1,-1),'LEFT'),
		('VALIGN',(0,0),(-1,-1),'MIDDLE'),
		('GRID',(0,0),(-1,-1),0.25,colors.black),
		])


    elements.append(TableHead)



    doc.build(elements)

    return buff







# Печатная форма для штампа согласования вариант формата docx
def Stamp2(d_id,persons):


    ### --- Список согласователей ---
    per = FIO_Job_Person(d_id,persons)

    ### --- Руководитель
    boss = FIO_Job_Boss(d_id)

    f = StringIO.StringIO()

    document = Document()


    ### Заголовок
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'КИС модуль согласования документов\nЗаяка № %s' % d_id).bold = True

    table = document.add_table(2+len(per),4, style="TableGrid")
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Pt(70)
    table.columns[0].width = Pt(150)
    table.columns[1].width = Pt(150)
    table.columns[2].width = Pt(70)    
    table.columns[3].width = Pt(120)    
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    heading_cells[0].text = u"Должность"
    heading_cells[1].text = u"ФИО"
    heading_cells[2].text = u"Дата"
    heading_cells[3].text = u"Подпись"    
    heading_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    heading_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    heading_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    heading_cells[3].paragraphs[0].runs[0].font.size = Pt(12)    
    heading_cells[0].paragraphs[0].runs[0].bold = True
    heading_cells[1].paragraphs[0].runs[0].bold = True
    heading_cells[2].paragraphs[0].runs[0].bold = True
    heading_cells[3].paragraphs[0].runs[0].bold = True    
    heading_cells[0].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[3].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="DCDCDC"/>'.format(nsdecls('w')))
    shading_elm_2 = parse_xml(r'<w:shd {} w:fill="DCDCDC"/>'.format(nsdecls('w')))
    shading_elm_3 = parse_xml(r'<w:shd {} w:fill="DCDCDC"/>'.format(nsdecls('w')))
    shading_elm_4 = parse_xml(r'<w:shd {} w:fill="DCDCDC"/>'.format(nsdecls('w')))
    table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
    table.rows[0].cells[1]._tc.get_or_add_tcPr().append(shading_elm_2)
    table.rows[0].cells[2]._tc.get_or_add_tcPr().append(shading_elm_3)
    table.rows[0].cells[3]._tc.get_or_add_tcPr().append(shading_elm_4)
    ### Руководитель
    cells = table.rows[1].cells
    cells[0].text = boss[0]
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].text = boss[1]
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    cells[2].text = GetSingD(d_id)
    cells[2].paragraphs[0].runs[0].font.size = Pt(10)

    ## Подпись руководителя - если есть
    if boss[2]:
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[1].height = Pt(70)
        parag = cells[3].paragraphs[0]
        run = parag.add_run()
        run.add_picture(StringIO.StringIO(boss[2]),width=Pt(100))


    n = 2
    for item in per:
        cells = table.rows[n].cells
        cells[0].text = item[0]
        cells[1].text = item[1]
        cells[2].text = GetSingP(d_id,item[2].encode('utf-8'))        
        cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        cells[2].paragraphs[0].runs[0].font.size = Pt(10)        
        if item[3]:
            table.rows[n].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            table.rows[n].height = Pt(70)
            parag = cells[3].paragraphs[0]
            run = parag.add_run()
            run.add_picture(StringIO.StringIO(item[3]),width=Pt(100))
        n += 1

    document.save(f)

    data = f.getvalue()

    return data
